import os
from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def getPlainFormattedDocument():
    doc = Document()
    section = doc.sections[0]
    section.right_margin = Mm(10)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    return doc


def createTable(docname, header, data):
    doc = Document()

    cols = len(header)

    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i in range(len(header)):
        hdr_cells[i].text = header[i]

    for i in range(len(data)):
        row_cells = table.add_row().cells
        for j in range(len(data[i])):
            row_cells[j].text = data[i][j]

    doc.save(docname)


def addPublication(outer_cntr, inner_cntr, table, publication):
    row = table.add_row().cells
    row[0].text = str(outer_cntr) + '.' + str(inner_cntr)
    row[1].text = publication.title
    row[2].text = str(len(publication.authors))
    if publication.scopusId != '-':
        row[5].text = '+'
    else:
        row[5].text = '-'
    if publication.ut != '-':
        row[5].text += '/+'
    else:
        row[5].text += '/-'
    row[7].text = 'DOI: ' + publication.doi

def createForm2(publications, path):
    doc = getPlainFormattedDocument()

    table = doc.add_table(rows=3, cols=8)
    table.style = 'Table Grid'

    # create complex header
    header_1 = ["№ п/п",
                "Опубликованные работы",
                "Численность\nработников\nкафедры, имеющих\nперечисленный\nрезультат",
                "Издание\nиз\nперечня\nВАК*",
                "Индексирование в\n  информационно-\nаналитических системах",
                "Ссылка в\nнаукометрической\nсистеме (DOI,\nРИНЦ, ФИПС,\nISSN)"]
    header_2 = ["РИНЦ*", "Scopus/Web of Science", "Другие"]

    for i in range(0, 4):
        cell1 = table.cell(0, i)
        cell2 = table.cell(1, i)
        cell3 = table.cell(2, i)
        cell1 = cell1.merge(cell2)
        cell1.text = header_1[i]
        cell3.text = str(i + 1)

    cell1 = table.cell(0, 4)
    cell2 = table.cell(0, 6)
    cell1 = cell1.merge(cell2)
    cell1.text = header_1[4]
    for i in range(4, 7):
        cell2 = table.cell(1, i)
        cell3 = table.cell(2, i)
        cell2.text = header_2[i - 4]
        cell3.text = str(i + 1)

    cell1 = table.cell(0, 7)
    cell2 = table.cell(1, 7)
    cell3 = table.cell(2, 7)
    cell1 = cell1.merge(cell2)
    cell1.text = header_1[5]
    cell3.text = '8'

    # fill articles
    cells = table.add_row().cells
    cells[0].text = '1'
    for i in range(2, len(cells)):
        cells[1].merge(cells[i])
    cells[1].text = "Статьи в научных журналах"
    counter = 1
    for publ in publications:
        if publ.type == "Article":
            addPublication(1, counter, table, publ)
            counter += 1

    # fill other publications
    cells = table.add_row().cells
    cells[0].text = '2'
    for i in range(2, len(cells)):
        cells[1].merge(cells[i])
    cells[1].text = "Нераспределенные публикации"
    counter = 1
    for publ in publications:
        if publ.type != "Article":
            addPublication(2, counter, table, publ)
            counter += 1
        # row = table.add_row().cells
        # row[0].text = str(i)
        # i += 1
        # row[1].text = publ.title
        # row[2].text = str(len(publ.authors))
        # if publ.scopusId != '-':
        #     row[5].text = '+'
        # else:
        #     row[5].text = '-'
        # if publ.ut != '-':
        #     row[5].text += '/+'
        # else:
        #     row[5].text += '/-'
        # row[7].text = 'DOI: ' + publ.doi

    if os.access(os.path.dirname(path), os.W_OK):
        doc.save(path)
        os.startfile(path)
